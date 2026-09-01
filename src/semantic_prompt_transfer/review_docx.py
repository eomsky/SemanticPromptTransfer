from __future__ import annotations

import io
import re
import zipfile
from html import escape
from pathlib import Path
from typing import Any, Iterable

from .domain import CaseContext, ReviewItem, ReviewSectionDraft
from .version import PACKAGE_VERSION


class OpinionDocumentBuilder:
    """Human-verifiable review document with numbered evidence appendix."""

    _citation = re.compile(r"\[?\s*((?:CR|ATT)_[a-f0-9]{20})\s*\]?", re.IGNORECASE)

    def __init__(self, capture_service=None) -> None:
        self.capture_service = capture_service
        self._bookmark_id = 1

    @classmethod
    def _visible_text(cls, text: str) -> str:
        value = cls._citation.sub("", str(text or ""))
        value = re.sub(r"\[\s*(?:,\s*)+\]", "", value)
        value = re.sub(r"\s+([.,;:])", r"\1", value)
        return re.sub(r"[ \t]{2,}", " ", value).strip()

    def _bookmark(self, paragraph, name: str) -> None:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        bookmark_id = str(self._bookmark_id)
        self._bookmark_id += 1
        start = OxmlElement("w:bookmarkStart")
        start.set(qn("w:id"), bookmark_id)
        start.set(qn("w:name"), name)
        end = OxmlElement("w:bookmarkEnd")
        end.set(qn("w:id"), bookmark_id)
        paragraph._p.insert(0, start)
        paragraph._p.append(end)

    @staticmethod
    def _anchor_link(paragraph, text: str, anchor: str) -> None:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("w:anchor"), anchor)
        run = OxmlElement("w:r")
        props = OxmlElement("w:rPr")
        color = OxmlElement("w:color")
        color.set(qn("w:val"), "2F5597")
        underline = OxmlElement("w:u")
        underline.set(qn("w:val"), "single")
        props.append(color)
        props.append(underline)
        text_node = OxmlElement("w:t")
        text_node.text = text
        run.append(props)
        run.append(text_node)
        hyperlink.append(run)
        paragraph._p.append(hyperlink)

    def _add_opinion_text(self, paragraph, text: str, refs: Iterable[dict[str, Any]]) -> None:
        mapping: dict[str, int] = {}
        for ref in refs:
            ref_no = int(ref.get("ref_no") or 0)
            for evidence_id in ref.get("member_ids", []):
                mapping[str(evidence_id)] = ref_no

        cursor = 0
        last_ref: int | None = None
        value = str(text or "")
        for match in self._citation.finditer(value):
            if match.start() > cursor:
                paragraph.add_run(value[cursor : match.start()])
            evidence_id = match.group(1)
            ref_no = mapping.get(evidence_id)
            if ref_no and ref_no != last_ref:
                self._anchor_link(paragraph, f"[근거 {ref_no}]", f"evidence_{ref_no}")
                last_ref = ref_no
            cursor = match.end()
        if cursor < len(value):
            paragraph.add_run(value[cursor:])
        if not paragraph.text.strip():
            paragraph.add_run(self._visible_text(text))

    @staticmethod
    def _all_refs(sections: Iterable[ReviewSectionDraft]) -> list[dict[str, Any]]:
        unique: dict[int, dict[str, Any]] = {}
        for section in sections:
            for ref in section.evidence_refs:
                ref_no = int(ref.get("ref_no") or 0)
                if ref_no and ref_no not in unique:
                    unique[ref_no] = dict(ref)
        return [unique[key] for key in sorted(unique)]

    @staticmethod
    def _fit_picture_dimensions(
        pixel_width: int,
        pixel_height: int,
        max_width_cm: float,
        max_height_cm: float,
    ) -> tuple[float, float]:
        width = max(1, int(pixel_width))
        height = max(1, int(pixel_height))
        target_width = float(max_width_cm)
        target_height = target_width * height / width
        if target_height > max_height_cm:
            target_height = float(max_height_cm)
            target_width = target_height * width / height
        return max(1.0, target_width), max(1.0, target_height)

    def _add_fitted_picture(self, document, png: bytes, section) -> tuple[float, float]:
        from PIL import Image
        from docx.shared import Cm

        with Image.open(io.BytesIO(png)) as image:
            pixel_width, pixel_height = image.size
        emu_per_cm = 360000.0
        max_width_cm = (section.page_width - section.left_margin - section.right_margin) / emu_per_cm
        usable_height_cm = (section.page_height - section.top_margin - section.bottom_margin) / emu_per_cm
        max_height_cm = max(8.0, usable_height_cm - 5.2)
        width_cm, height_cm = self._fit_picture_dimensions(
            pixel_width, pixel_height, max_width_cm, max_height_cm
        )
        document.add_picture(io.BytesIO(png), width=Cm(width_cm), height=Cm(height_cm))
        return width_cm, height_cm

    @staticmethod
    def _safe_evidence_excerpt(value: str) -> str:
        cleaned = re.sub(r"\[(?:VALUE|PERIOD|COMPANY|DATE|ENTITY)\]", "", str(value or ""), flags=re.I)
        return cleaned.strip() or "원문 캡처를 생성하지 못했습니다."

    def build(
        self,
        case: CaseContext,
        sections: Iterable[ReviewSectionDraft],
        output_path: str | Path,
        *,
        title: str = "여신 심사의견",
        evidence_catalog: dict[str, dict[str, Any]] | None = None,
    ) -> Path:
        try:
            from docx import Document
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            from docx.shared import Cm, Pt
        except ImportError as exc:
            raise RuntimeError("python-docx is required for Word output") from exc

        rows = list(sections)
        by_item = {section.review_item: section for section in rows}
        missing = [item.value for item in ReviewItem.ordered() if item not in by_item]
        if missing:
            raise ValueError(f"review sections missing: {missing}")
        evidence_catalog = dict(evidence_catalog or {})

        document = Document()
        section = document.sections[0]
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

        for name in ("Normal", "Title", "Heading 1", "Heading 2"):
            style = document.styles[name]
            style.font.name = "Arial"
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
            if name == "Normal":
                style.font.size = Pt(10.5)

        title_p = document.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_p.add_run(title)
        title_run.bold = True
        title_run.font.size = Pt(21)
        self._bookmark(title_p, "review_top")

        # Pending future business-metadata redesign: show only the case id.
        info = document.add_table(rows=1, cols=2)
        info.style = "Table Grid"
        info.rows[0].cells[0].text = "심사건"
        info.rows[0].cells[1].text = case.case_id
        info.rows[0].cells[0].paragraphs[0].runs[0].bold = True
        document.add_paragraph()

        korean_labels = ("가", "나", "다", "라", "마")
        for label, item in zip(korean_labels, ReviewItem.ordered(), strict=True):
            current = by_item[item]
            heading = document.add_paragraph(style="Heading 1")
            heading.paragraph_format.space_before = Pt(12)
            heading.paragraph_format.space_after = Pt(4)
            heading.add_run(f"{label}. {item.title}").bold = True

            body = document.add_paragraph()
            body.paragraph_format.line_spacing = 1.25
            body.paragraph_format.space_after = Pt(6)
            self._add_opinion_text(body, current.text, current.evidence_refs)

            if current.evidence_refs:
                refs = ", ".join(f"근거 {int(ref['ref_no'])}" for ref in current.evidence_refs)
                note = document.add_paragraph()
                note.paragraph_format.space_after = Pt(5)
                run = note.add_run(f"참조: {refs}")
                run.italic = True
                run.font.size = Pt(8.5)

        refs = self._all_refs(rows)
        if refs:
            document.add_page_break()
            appendix = document.add_paragraph(style="Heading 1")
            appendix.add_run("근거 자료").bold = True
            intro = document.add_paragraph(
                "심사의견에 표시된 [근거 N] 번호와 아래 근거 이미지를 대조하여 원문을 직접 확인할 수 있습니다."
            )
            intro.paragraph_format.space_after = Pt(10)
            document.add_page_break()

            for index, ref in enumerate(refs):
                if index:
                    document.add_page_break()
                ref_no = int(ref["ref_no"])
                heading = document.add_paragraph(style="Heading 2")
                heading.add_run(f"근거 {ref_no} · {ref.get('source_filename') or '업로드 자료'}").bold = True
                self._bookmark(heading, f"evidence_{ref_no}")

                meta = document.add_paragraph()
                meta.add_run(f"위치: {ref.get('location') or '-'}").bold = True
                used_by = ", ".join(ref.get("review_items", []))
                if used_by:
                    meta.add_run(f"   |   활용 항목: {used_by}")

                representative_id = str(ref.get("representative_id") or "")
                source_row = dict(evidence_catalog.get(representative_id) or {})
                metadata = dict(source_row.get("metadata") or {})
                if ref.get("highlight_bboxes"):
                    metadata["highlight_bboxes"] = ref["highlight_bboxes"]
                if ref.get("highlight_cell_ranges"):
                    metadata["highlight_cell_ranges"] = ref["highlight_cell_ranges"]
                source_row["metadata"] = metadata

                rendered = False
                if self.capture_service is not None and source_row:
                    try:
                        png = self.capture_service.capture_png(case.tenant_id, case.case_id, source_row)
                        self._add_fitted_picture(document, png, section)
                        rendered = True
                    except Exception:
                        rendered = False
                if not rendered:
                    excerpt = self._safe_evidence_excerpt(str(source_row.get("content") or ""))
                    document.add_paragraph(excerpt[:1800])

                back = document.add_paragraph()
                self._anchor_link(back, "↩ 심사의견으로 돌아가기", "review_top")

        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.add_run(f"SemanticPromptTransfer v{PACKAGE_VERSION} · 근거 추적형 생성문서")

        settings = document.settings.element
        update_fields = OxmlElement("w:updateFields")
        update_fields.set(qn("w:val"), "true")
        settings.append(update_fields)

        target = Path(output_path)
        target.parent.mkdir(parents=True, exist_ok=True)
        document.save(target)
        return target

    def build_minimal(
        self,
        case: CaseContext,
        sections: Iterable[ReviewSectionDraft],
        output_path: str | Path,
        *,
        title: str = "여신 심사의견",
    ) -> Path:
        rows = list(sections)
        by_item = {section.review_item: section for section in rows}
        paragraphs = [title, f"심사건: {case.case_id}"]
        for item in ReviewItem.ordered():
            paragraphs.append(f"{item.value}. {item.title}")
            paragraphs.append(
                self._visible_text(by_item[item].text) if item in by_item else "문구를 확정하지 못했습니다."
            )
        body = "".join(
            f'<w:p><w:r><w:t xml:space="preserve">{escape(text)}</w:t></w:r></w:p>'
            for text in paragraphs
        )
        document_xml = (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            f'<w:body>{body}<w:sectPr/></w:body></w:document>'
        )
        content_types = (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
            '<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>'
            '<Default Extension="xml" ContentType="application/xml"/>'
            '<Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>'
            '</Types>'
        )
        rels = (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
            '<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>'
            '</Relationships>'
        )
        target = Path(output_path)
        target.parent.mkdir(parents=True, exist_ok=True)
        with zipfile.ZipFile(target, "w", compression=zipfile.ZIP_DEFLATED) as archive:
            archive.writestr("[Content_Types].xml", content_types)
            archive.writestr("_rels/.rels", rels)
            archive.writestr("word/document.xml", document_xml)
        return target
