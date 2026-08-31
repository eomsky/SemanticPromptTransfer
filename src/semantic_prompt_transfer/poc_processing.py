from __future__ import annotations

import json
import hashlib
import re
import threading
from collections import OrderedDict, Counter
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Callable

import numpy as np
from docx import Document
from openpyxl import load_workbook

from ._chunk_builder_base import ChunkRecord
from .config import DocumentScope
from .credit_report import CreditReportParser, CreditReportTemplate
from .domain import DocumentKind, FileStatus
from .encoding import EncoderBackend
from .storage import LocalDocumentArtifactStore
from .vector_store import ShardedNpzVectorStore


@dataclass(frozen=True)
class ExtractedBlock:
    text: str
    page: int | None = None
    location: str | None = None
    bbox: tuple[float, float, float, float] | None = None
    page_size: tuple[float, float] | None = None
    source_spans: tuple[dict[str, Any], ...] = ()


class PocDocumentExtractor:
    """Small deterministic extractor for the temporary POC upload path."""

    SUPPORTED = {".pdf", ".docx", ".xlsx", ".txt", ".md"}

    def extract(self, path: str | Path) -> list[ExtractedBlock]:
        source = Path(path)
        suffix = source.suffix.lower()
        if suffix not in self.SUPPORTED:
            raise ValueError(f"unsupported attachment type: {suffix or 'unknown'}")
        if suffix == ".pdf":
            return self._pdf(source)
        if suffix == ".docx":
            return self._docx(source)
        if suffix == ".xlsx":
            return self._xlsx(source)
        return self._text(source)

    @staticmethod
    def _pdf(path: Path) -> list[ExtractedBlock]:
        try:
            import fitz
        except ImportError as exc:  # pragma: no cover - optional POC dependency
            raise RuntimeError("install semantic-prompt-transfer[poc] for PDF coordinate extraction") from exc
        document = fitz.open(str(path))
        blocks: list[ExtractedBlock] = []
        try:
            for page_index, page in enumerate(document, start=1):
                page_size = (float(page.rect.width), float(page.rect.height))
                for block in page.get_text("blocks", sort=True):
                    x0, y0, x1, y1, text = block[:5]
                    normalized = " ".join(str(text or "").split())
                    if not normalized:
                        continue
                    blocks.append(
                        ExtractedBlock(
                            text=normalized,
                            page=page_index,
                            location=f"page:{page_index}:block:{len(blocks) + 1}",
                            bbox=(float(x0), float(y0), float(x1), float(y1)),
                            page_size=page_size,
                        )
                    )
        finally:
            document.close()
        return blocks

    @staticmethod
    def _docx(path: Path) -> list[ExtractedBlock]:
        document = Document(str(path))
        blocks = [
            ExtractedBlock(text=paragraph.text.strip(), location=f"paragraph:{index}")
            for index, paragraph in enumerate(document.paragraphs, start=1)
            if paragraph.text.strip()
        ]
        for table_index, table in enumerate(document.tables, start=1):
            for row_index, row in enumerate(table.rows, start=1):
                values = [" ".join(cell.text.split()) for cell in row.cells]
                if any(values):
                    blocks.append(
                        ExtractedBlock(
                            text=" | ".join(values),
                            location=f"table:{table_index}:row:{row_index}",
                        )
                    )
        return blocks

    @staticmethod
    def _xlsx(path: Path) -> list[ExtractedBlock]:
        workbook = load_workbook(path, data_only=False, read_only=True)
        blocks: list[ExtractedBlock] = []
        try:
            for sheet in workbook.worksheets:
                for row_index, row in enumerate(sheet.iter_rows(values_only=False), start=1):
                    values = []
                    for cell in row:
                        if cell.value is not None:
                            values.append(f"{cell.coordinate}={cell.value}")
                    if values:
                        blocks.append(
                            ExtractedBlock(
                                text=f"시트={sheet.title}; " + "; ".join(values),
                                location=f"sheet:{sheet.title}:row:{row_index}",
                            )
                        )
        finally:
            workbook.close()
        return blocks

    @staticmethod
    def _text(path: Path) -> list[ExtractedBlock]:
        payload = path.read_bytes()
        text = None
        for encoding in ("utf-8-sig", "cp949", "utf-8"):
            try:
                text = payload.decode(encoding)
                break
            except UnicodeDecodeError:
                continue
        if text is None:
            raise ValueError("text attachment encoding is not supported")
        return [ExtractedBlock(text=text.strip(), location="text")]


class PocUploadProcessor:
    """Parse, chunk and embed uploads entirely inside one ephemeral runtime."""

    def __init__(
        self,
        encoder: EncoderBackend,
        vector_store: ShardedNpzVectorStore,
        artifact_store: LocalDocumentArtifactStore,
        *,
        credit_template: CreditReportTemplate | None = None,
        extractor: PocDocumentExtractor | None = None,
        max_chars: int = 1000,
        overlap_chars: int = 100,
        embedding_cache_size: int = 8,
    ) -> None:
        if max_chars < 256 or not 0 <= overlap_chars < max_chars:
            raise ValueError("invalid POC chunk size or overlap")
        self.encoder = encoder
        self.vector_store = vector_store
        self.artifact_store = artifact_store
        self.credit_template = credit_template
        self.extractor = extractor or PocDocumentExtractor()
        self.max_chars = int(max_chars)
        self.overlap_chars = int(overlap_chars)
        self.embedding_cache_size = max(0, int(embedding_cache_size))
        self._embedding_cache: OrderedDict[
            tuple[str, str, str, str], tuple[tuple[str, ...], np.ndarray]
        ] = OrderedDict()
        self._cache_lock = threading.RLock()

    @staticmethod
    def _write_json(path: Path, value: Any) -> None:
        path.parent.mkdir(parents=True, exist_ok=True)
        temporary = path.with_suffix(path.suffix + ".tmp")
        temporary.write_text(json.dumps(value, ensure_ascii=False, indent=2), encoding="utf-8")
        temporary.replace(path)

    def _split(self, text: str) -> list[str]:
        normalized = re.sub(r"[ \t]+", " ", text).strip()
        if not normalized:
            return []
        chunks: list[str] = []
        start = 0
        while start < len(normalized):
            end = min(len(normalized), start + self.max_chars)
            if end < len(normalized):
                boundary = max(
                    normalized.rfind("\n", start + self.max_chars // 2, end),
                    normalized.rfind(". ", start + self.max_chars // 2, end),
                    normalized.rfind(" ", start + self.max_chars // 2, end),
                )
                if boundary > start:
                    end = boundary + 1
            chunk = normalized[start:end].strip()
            if chunk:
                chunks.append(chunk)
            if end >= len(normalized):
                break
            next_start = max(start + 1, end - self.overlap_chars)
            start = next_start
        return chunks

    @staticmethod
    def _source_span(block: ExtractedBlock) -> dict[str, Any]:
        return {
            "text": block.text,
            "page": block.page,
            "location": block.location,
            "bbox": list(block.bbox) if block.bbox else None,
            "page_size": list(block.page_size) if block.page_size else None,
        }

    @staticmethod
    def _is_margin_block(block: ExtractedBlock) -> bool:
        if not block.bbox or not block.page_size:
            return False
        _, y0, _, y1 = block.bbox
        height = max(float(block.page_size[1]), 1.0)
        return y1 <= height * 0.08 or y0 >= height * 0.92

    def _filter_repeated_margins(self, blocks: list[ExtractedBlock]) -> list[ExtractedBlock]:
        counts = Counter(
            " ".join(block.text.lower().split())
            for block in blocks
            if block.page and self._is_margin_block(block) and len(block.text.strip()) <= 160
        )
        repeated = {text for text, count in counts.items() if text and count >= 3}
        if not repeated:
            return blocks
        return [
            block
            for block in blocks
            if not (
                self._is_margin_block(block)
                and " ".join(block.text.lower().split()) in repeated
            )
        ]

    def _coalesce_pdf_blocks(self, blocks: list[ExtractedBlock]) -> list[ExtractedBlock]:
        """Merge only nearby blocks on one page and retain every source span."""
        prepared = self._filter_repeated_margins(blocks)
        merged: list[ExtractedBlock] = []
        current: list[ExtractedBlock] = []

        def flush() -> None:
            if not current:
                return
            if len(current) == 1:
                block = current[0]
                spans = block.source_spans or (self._source_span(block),)
                merged.append(
                    ExtractedBlock(
                        block.text,
                        block.page,
                        block.location,
                        block.bbox,
                        block.page_size,
                        tuple(spans),
                    )
                )
                current.clear()
                return
            boxes = [block.bbox for block in current if block.bbox]
            bbox = (
                min(box[0] for box in boxes),
                min(box[1] for box in boxes),
                max(box[2] for box in boxes),
                max(box[3] for box in boxes),
            ) if boxes else None
            spans = tuple(
                span
                for block in current
                for span in (block.source_spans or (self._source_span(block),))
            )
            merged.append(
                ExtractedBlock(
                    text="\n".join(block.text for block in current),
                    page=current[0].page,
                    location=f"{current[0].location}..{current[-1].location}",
                    bbox=bbox,
                    page_size=current[0].page_size,
                    source_spans=spans,
                )
            )
            current.clear()

        for block in prepared:
            if not current:
                current.append(block)
                continue
            previous = current[-1]
            same_page = block.page is not None and block.page == previous.page
            combined_chars = sum(len(value.text) for value in current) + len(block.text)
            nearby = False
            if previous.bbox and block.bbox:
                vertical_gap = float(block.bbox[1]) - float(previous.bbox[3])
                horizontal_overlap = min(previous.bbox[2], block.bbox[2]) - max(
                    previous.bbox[0], block.bbox[0]
                )
                nearby = vertical_gap <= 28.0 and horizontal_overlap >= -12.0
            if same_page and nearby and combined_chars <= self.max_chars:
                current.append(block)
            else:
                flush()
                current.append(block)
        flush()
        return merged

    @staticmethod
    def _file_sha256(path: Path) -> str:
        digest = hashlib.sha256()
        with path.open("rb") as stream:
            for block in iter(lambda: stream.read(1024 * 1024), b""):
                digest.update(block)
        return digest.hexdigest()

    def _cache_key(self, scope: DocumentScope, source_path: Path) -> tuple[str, str, str, str]:
        metadata = self.encoder.metadata()
        profile = "|".join(
            str(metadata.get(key) or "")
            for key in ("model_id", "provider", "max_length", "dimension")
        )
        return scope.tenant_id, scope.case_id, self._file_sha256(source_path), profile

    def _cached_embeddings(
        self,
        key: tuple[str, str, str, str],
        texts: list[str],
    ) -> np.ndarray | None:
        signatures = tuple(hashlib.sha256(text.encode("utf-8")).hexdigest() for text in texts)
        with self._cache_lock:
            cached = self._embedding_cache.get(key)
            if cached is None or cached[0] != signatures:
                return None
            self._embedding_cache.move_to_end(key)
            return cached[1].copy()

    def _remember_embeddings(
        self,
        key: tuple[str, str, str, str],
        texts: list[str],
        embeddings: np.ndarray,
    ) -> None:
        if self.embedding_cache_size < 1:
            return
        signatures = tuple(hashlib.sha256(text.encode("utf-8")).hexdigest() for text in texts)
        with self._cache_lock:
            self._embedding_cache[key] = (
                signatures,
                np.asarray(embeddings, dtype=np.float32).copy(),
            )
            self._embedding_cache.move_to_end(key)
            while len(self._embedding_cache) > self.embedding_cache_size:
                self._embedding_cache.popitem(last=False)

    def _records(
        self,
        scope: DocumentScope,
        source_path: Path,
        blocks: list[ExtractedBlock],
    ) -> list[ChunkRecord]:
        records: list[ChunkRecord] = []
        source_blocks = (
            self._coalesce_pdf_blocks(blocks)
            if source_path.suffix.lower() == ".pdf"
            else blocks
        )
        for block_index, block in enumerate(source_blocks, start=1):
            for part_index, text in enumerate(self._split(block.text), start=1):
                local_id = f"CH_POC_{len(records) + 1:05d}"
                global_id = DocumentScope(
                    scope.tenant_id, scope.case_id, scope.document_id
                ).as_metadata()
                identity = "\x1f".join(
                    (scope.tenant_id, scope.case_id, scope.document_id, local_id, "0")
                )
                chunk_hash = hashlib.sha256(identity.encode("utf-8")).hexdigest()
                metadata = {
                    **global_id,
                    "global_chunk_id": f"GCH_{chunk_hash}",
                    "local_chunk_id": local_id,
                    "representation_level": 0,
                    "source_filename": source_path.name,
                    "document_kind": DocumentKind.ATTACHMENT.value,
                    "pages": [block.page] if block.page else [],
                    "source_location": block.location,
                    "bbox": list(block.bbox) if block.bbox else None,
                    "page_size": list(block.page_size) if block.page_size else None,
                    "source_spans": list(block.source_spans),
                    "block_index": block_index,
                    "part_index": part_index,
                }
                records.append(
                    ChunkRecord(
                        chunk_id=local_id,
                        embedding_text=text,
                        document=text,
                        metadata=metadata,
                    )
                )
        return records

    def process(
        self,
        scope: DocumentScope,
        source_path: Path,
        document_kind: DocumentKind,
        progress: Callable[[FileStatus, int | None, str | None], None],
    ) -> None:
        derived = self.artifact_store.derived_path(scope)
        progress(FileStatus.VALIDATING, 25, "파일 형식과 범위를 검증했습니다.")
        if document_kind is DocumentKind.CREDIT_REPORT:
            if source_path.suffix.lower() != ".xlsx":
                raise ValueError("POC 신용조사서는 .xlsx 형식이어야 합니다")
            if self.credit_template is None:
                raise RuntimeError("credit-report template is not configured")
            progress(FileStatus.PARSING, 55, "신용조사서 정형 셀을 읽고 있습니다.")
            parser_scope = DocumentScope(
                scope.tenant_id,
                scope.case_id,
                scope.document_id,
                source_filename=source_path.name,
                document_kind=DocumentKind.CREDIT_REPORT.value,
            )
            parser = CreditReportParser()
            try:
                parsed = parser.parse_semantic_workbook(source_path, parser_scope)
            except ValueError as exc:
                if "missing operational credit-report sheets" not in str(exc):
                    raise
                parsed = parser.parse(source_path, self.credit_template, parser_scope)
            self._write_json(derived / "credit_facts.json", parsed.to_dict())
            self.vector_store.delete_document(scope)
            progress(FileStatus.READY, 100, "신용조사서 기초자료 적재를 완료했습니다.")
            return

        progress(FileStatus.PARSING, 45, "첨부자료 텍스트를 추출하고 있습니다.")
        blocks = self.extractor.extract(source_path)
        progress(FileStatus.PARSING, 60, "본문 블록을 검색 청크로 병합하고 있습니다.")
        records = self._records(scope, source_path, blocks)
        if not records:
            raise ValueError("첨부자료에서 검색 가능한 텍스트를 찾지 못했습니다")
        texts = [row.embedding_text for row in records]
        cache_key = self._cache_key(scope, source_path)
        embeddings = self._cached_embeddings(cache_key, texts)
        if embeddings is None:
            progress(FileStatus.INDEXING, 70, f"GPU 임베딩 준비 · {len(texts):,}개 청크")
            accelerated = getattr(self.encoder, "encode_documents_with_progress", None)
            if callable(accelerated):
                def on_batch(done: int, total: int) -> None:
                    percent = 72 + int(20 * done / max(total, 1))
                    progress(
                        FileStatus.INDEXING,
                        min(percent, 92),
                        f"GPU 임베딩 배치 {done:,}/{total:,}",
                    )

                embeddings = np.asarray(accelerated(texts, on_batch), dtype=np.float32)
            else:
                embeddings = np.asarray(self.encoder.encode_documents(texts), dtype=np.float32)
                progress(FileStatus.INDEXING, 92, "임베딩 벡터 생성을 완료했습니다.")
            self._remember_embeddings(cache_key, texts, embeddings)
        else:
            progress(FileStatus.INDEXING, 92, "동일 파일의 임베딩 캐시를 재사용했습니다.")
        progress(FileStatus.INDEXING, 95, "벡터 인덱스를 저장하고 있습니다.")
        self.vector_store.upsert_document(records, embeddings)
        progress(FileStatus.INDEXING, 98, "근거 좌표와 청크 메타데이터를 저장하고 있습니다.")
        self._write_json(
            derived / "chunks.json",
            {
                "document_id": scope.document_id,
                "source_filename": source_path.name,
                "chunk_count": len(records),
                "chunks": [
                    {
                        "chunk_id": row.chunk_id,
                        "embedding_text": row.embedding_text,
                        "document": row.document,
                        "metadata": row.metadata,
                    }
                    for row in records
                ],
            },
        )
        progress(FileStatus.READY, 100, "첨부자료 임베딩과 벡터 적재를 완료했습니다.")


class ShardedAttachmentRetriever:
    def __init__(
        self,
        encoder: EncoderBackend,
        vector_store: ShardedNpzVectorStore,
        *,
        top_k: int = 8,
    ) -> None:
        self.encoder = encoder
        self.vector_store = vector_store
        self.top_k = int(top_k)

    @staticmethod
    def _dedupe_key(hit: dict[str, Any]) -> str:
        text = " ".join(str(hit.get("document") or hit.get("embedding_text") or "").lower().split())
        return hashlib.sha256(text[:1200].encode("utf-8")).hexdigest()

    def search(self, query: str, **kwargs: Any) -> dict[str, Any]:
        filters = dict(kwargs.get("filters") or {})
        missing = [key for key in ("tenant_id", "case_id") if not filters.get(key)]
        if missing:
            raise ValueError(f"POC retrieval requires scope filters: {missing}")
        embedding = self.encoder.encode_queries([query])[0]
        raw = self.vector_store.search(embedding, top_k=max(self.top_k * 4, self.top_k), filters=filters)
        if not raw:
            return {"query": query, "filters": filters, "hits": [], "retrieval_quality": "empty"}

        best = float(raw[0].get("score") or 0.0)
        # Dynamic relevance gate: retain results close to the best match, but do not
        # force unrelated evidence into the prompt when every cosine score is weak.
        floor = max(0.10, best - 0.12)
        accepted: list[dict[str, Any]] = []
        seen: set[str] = set()
        for hit in raw:
            score = float(hit.get("score") or 0.0)
            if score < floor:
                continue
            key = self._dedupe_key(hit)
            if key in seen:
                continue
            seen.add(key)
            row = dict(hit)
            row["relevance_status"] = "accepted"
            accepted.append(row)
            if len(accepted) >= self.top_k:
                break
        quality = "accepted" if accepted else "low_relevance"
        return {
            "query": query,
            "filters": filters,
            "hits": accepted,
            "retrieval_quality": quality,
            "best_score": best,
            "score_floor": floor,
            "raw_hit_count": len(raw),
        }
