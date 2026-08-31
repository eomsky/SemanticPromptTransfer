from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor


FONT = "NanumGothic"
BLUE = RGBColor(0x17, 0x36, 0x5D)
TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def _set_twips(parent, tag, value):
    node = parent.find(qn(tag))
    if node is None:
        node = OxmlElement(tag)
        parent.append(node)
    node.set(qn("w:w"), str(value))
    node.set(qn("w:type"), "dxa")


def set_table_geometry(table):
    columns = len(table.columns)
    patterns = {
        1: [TABLE_WIDTH_DXA],
        2: [3000, 6360],
        3: [2600, 3000, 3760],
        4: [1700, 2200, 2500, 2960],
    }
    widths = patterns.get(columns)
    if widths is None:
        base = TABLE_WIDTH_DXA // columns
        widths = [base] * columns
        widths[-1] += TABLE_WIDTH_DXA - sum(widths)

    table.autofit = False
    properties = table._tbl.tblPr
    _set_twips(properties, "w:tblW", TABLE_WIDTH_DXA)
    _set_twips(properties, "w:tblInd", TABLE_INDENT_DXA)
    layout = properties.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        properties.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row_index, row in enumerate(table.rows):
        if row_index == 0:
            row_properties = row._tr.get_or_add_trPr()
            if row_properties.find(qn("w:tblHeader")) is None:
                row_properties.append(OxmlElement("w:tblHeader"))
        for column_index, cell in enumerate(row.cells):
            cell_properties = cell._tc.get_or_add_tcPr()
            _set_twips(cell_properties, "w:tcW", widths[min(column_index, columns - 1)])
            margins = cell_properties.find(qn("w:tcMar"))
            if margins is None:
                margins = OxmlElement("w:tcMar")
                cell_properties.append(margins)
            for side, value in (("top", 80), ("bottom", 80), ("start", 120), ("end", 120)):
                node = margins.find(qn(f"w:{side}"))
                if node is None:
                    node = OxmlElement(f"w:{side}")
                    margins.append(node)
                node.set(qn("w:w"), str(value))
                node.set(qn("w:type"), "dxa")


def set_run_font(run, name=FONT, size=None, bold=None, color=None):
    run.font.name = name
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    fonts.set(qn("w:eastAsia"), name)
    fonts.set(qn("w:ascii"), name)
    fonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def add_page_field(paragraph):
    run = paragraph.add_run()
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    run._r.addnext(field)


def style_document(path: Path) -> None:
    document = Document(path)
    document.settings.odd_and_even_pages_header_footer = False
    for section in document.sections:
        section.page_width, section.page_height = Mm(210), Mm(297)
        section.top_margin = section.bottom_margin = Mm(22)
        section.left_margin = section.right_margin = Mm(22)
        section.header_distance = section.footer_distance = Mm(10)
        for header_part in (section.header, section.even_page_header, section.first_page_header):
            header = header_part.paragraphs[0]
            header.text = "SemanticPromptTransfer v0.21 | 운영 인터페이스 및 CPU 생성 요구사항"
            header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for run in header.runs:
                set_run_font(run, size=8.5, color=RGBColor(0x66, 0x66, 0x66))
        footer = section.footer.paragraphs[0]
        footer.text = "페이지 "
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in footer.runs:
            set_run_font(run, size=8.5, color=RGBColor(0x66, 0x66, 0x66))
        add_page_field(footer)

    normal = document.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.22

    styles = {
        "Title": (21, True, BLUE, 12),
        "Subtitle": (12, False, RGBColor(0x55, 0x55, 0x55), 10),
        "Heading 1": (16, True, BLUE, 8),
        "Heading 2": (13, True, BLUE, 6),
        "Heading 3": (11.5, True, RGBColor(0x1F, 0x4D, 0x78), 4),
    }
    for name, (size, bold, color, after) in styles.items():
        if name not in document.styles:
            continue
        style = document.styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.color.rgb = color
        style.paragraph_format.space_after = Pt(after)

    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            set_run_font(run)
    for table in document.tables:
        set_table_geometry(table)
        for row_index, row in enumerate(table.rows):
            for cell in row.cells:
                properties = cell._tc.get_or_add_tcPr()
                shade = properties.find(qn("w:shd"))
                if row_index == 0:
                    if shade is None:
                        shade = OxmlElement("w:shd")
                        properties.append(shade)
                    shade.set(qn("w:fill"), "E8EEF5")
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_after = Pt(2)
                    for run in paragraph.runs:
                        set_run_font(run, size=9.2, bold=True if row_index == 0 else None)

    document.core_properties.title = "SemanticPromptTransfer v0.21 운영 인터페이스 및 CPU 생성 요구사항"
    document.core_properties.subject = "여신심사 HTML, 파일 생명주기, 다중 문서 RAG, CPU 생성기 계약"
    document.core_properties.author = "SemanticPromptTransfer"
    document.save(path)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("path", type=Path)
    args = parser.parse_args()
    style_document(args.path.resolve())
    print(args.path.resolve())


if __name__ == "__main__":
    main()
