from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


FONT = "NanumGothic"
BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
MUTED = RGBColor(0x66, 0x66, 0x66)
TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def _set_twips(parent, tag, value):
    node = parent.find(qn(tag))
    if node is None:
        node = OxmlElement(tag)
        parent.append(node)
    node.set(qn("w:w"), str(value))
    node.set(qn("w:type"), "dxa")


def set_table_geometry(table):
    columns = len(table.columns)
    patterns = {
        1: [TABLE_WIDTH_DXA],
        2: [2700, 6660],
        3: [2500, 1600, 5260],
        4: [1700, 2000, 2460, 3200],
    }
    widths = patterns.get(columns)
    if widths is None:
        base = TABLE_WIDTH_DXA // columns
        widths = [base] * columns
        widths[-1] += TABLE_WIDTH_DXA - sum(widths)

    table.autofit = False
    properties = table._tbl.tblPr
    _set_twips(properties, "w:tblW", TABLE_WIDTH_DXA)
    _set_twips(properties, "w:tblInd", TABLE_INDENT_DXA)
    layout = properties.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        properties.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row_index, row in enumerate(table.rows):
        if row_index == 0:
            row_properties = row._tr.get_or_add_trPr()
            if row_properties.find(qn("w:tblHeader")) is None:
                row_properties.append(OxmlElement("w:tblHeader"))
        for column_index, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell_properties = cell._tc.get_or_add_tcPr()
            _set_twips(cell_properties, "w:tcW", widths[column_index])
            margins = cell_properties.find(qn("w:tcMar"))
            if margins is None:
                margins = OxmlElement("w:tcMar")
                cell_properties.append(margins)
            for side, value in (("top", 80), ("bottom", 80), ("start", 120), ("end", 120)):
                node = margins.find(qn(f"w:{side}"))
                if node is None:
                    node = OxmlElement(f"w:{side}")
                    margins.append(node)
                node.set(qn("w:w"), str(value))
                node.set(qn("w:type"), "dxa")


def set_run_font(run, name=FONT, size=None, bold=None, color=None):
    run.font.name = name
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    for key in ("w:eastAsia", "w:ascii", "w:hAnsi"):
        fonts.set(qn(key), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def add_page_field(paragraph):
    run = paragraph.add_run()
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    run._r.addnext(field)


def style_document(path: Path) -> None:
    document = Document(path)
    document.settings.odd_and_even_pages_header_footer = False
    for section in document.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = section.bottom_margin = Inches(1)
        section.left_margin = section.right_margin = Inches(1)
        section.header_distance = section.footer_distance = Inches(0.492)
        header = section.header.paragraphs[0]
        header.text = "SemanticPromptTransfer v0.22 | Colab POC 운영 요구사항"
        header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        header.paragraph_format.space_after = Pt(0)
        for run in header.runs:
            set_run_font(run, size=8.5, color=MUTED)
        footer = section.footer.paragraphs[0]
        footer.text = "페이지 "
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.paragraph_format.space_before = Pt(0)
        for run in footer.runs:
            set_run_font(run, size=8.5, color=MUTED)
        add_page_field(footer)

    normal = document.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    style_tokens = {
        "Title": (20, True, RGBColor(0, 0, 0), 0, 8),
        "Subtitle": (11, False, MUTED, 0, 12),
        "Heading 1": (16, True, BLUE, 16, 8),
        "Heading 2": (13, True, BLUE, 12, 6),
        "Heading 3": (12, True, DARK_BLUE, 8, 4),
    }
    for name, (size, bold, color, before, after) in style_tokens.items():
        if name not in document.styles:
            continue
        style = document.styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for index, paragraph in enumerate(document.paragraphs):
        if index == 0 and paragraph.text.strip().startswith("SemanticPromptTransfer v0.22"):
            paragraph.style = document.styles["Title"]
        for run in paragraph.runs:
            set_run_font(run)

    for table in document.tables:
        set_table_geometry(table)
        for row_index, row in enumerate(table.rows):
            for column_index, cell in enumerate(row.cells):
                properties = cell._tc.get_or_add_tcPr()
                if row_index == 0:
                    shade = properties.find(qn("w:shd"))
                    if shade is None:
                        shade = OxmlElement("w:shd")
                        properties.append(shade)
                    shade.set(qn("w:fill"), "F2F4F7")
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(2)
                    paragraph.paragraph_format.line_spacing = 1.10
                    if column_index == 1 and len(row.cells) >= 3:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        set_run_font(run, size=9.3, bold=True if row_index == 0 else None)

    document.core_properties.title = "SemanticPromptTransfer v0.22 Colab POC 운영 요구사항"
    document.core_properties.subject = "회원가입, 양식 다운로드, 다중 PDF RAG, LLM 교체, 검증 삭제"
    document.core_properties.author = "SemanticPromptTransfer"
    document.save(path)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("path", type=Path)
    args = parser.parse_args()
    style_document(args.path.resolve())
    print(args.path.resolve())


if __name__ == "__main__":
    main()
